
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_quotation():
    doc = Document()
    
    # Estilo general
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Encabezado / Título
    title = doc.add_heading('COTIZACIÓN DE SERVICIOS - DESARROLLO WEB', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Información de las partes
    doc.add_paragraph(f"FECHA: 4 de Febrero de 2026")
    doc.add_paragraph(f"CLIENTE: Dr. CESAR MIGUEL RIASCOS NOGUERA")
    doc.add_paragraph(f"DESARROLLADOR: JORGE JIMENEZ CALVANO")
    doc.add_paragraph("-" * 30)

    # Introducción
    doc.add_heading('1. DESCRIPCIÓN DEL PROYECTO', level=1)
    p = doc.add_paragraph()
    p.add_run('Desarrollo y diseño de la plataforma web ARTETEX, especializada en servicios de sublimación y estampado textil premium. La web ha sido construida bajo una estética minimalista suiza, priorizando la visualización de alta calidad y la experiencia de usuario.').bold = False

    # Características
    doc.add_heading('2. CARACTERÍSTICAS TÉCNICAS', level=1)
    features = [
        "Arquitectura Mobile-First optimizada para dispositivos móviles.",
        "Diseño Premium con estética suiza (Líneas limpias, tipografía imponente).",
        "Galería de Productos Curada con filtros de visualización.",
        "Laboratorio de IA integrado para prototipado de patrones técnicos.",
        "Optimización SEO para posicionamiento en motores de búsqueda.",
        "Sección Legal y Privacidad (Habeas Data Colombia).",
        "Integración con redes sociales (Instagram, LinkedIn, Behance).",
        "Efectos visuales dinámicos (Hover interactivo a color)."
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # Entrega
    doc.add_heading('3. ENTREGABLES', level=1)
    doc.add_paragraph("Código fuente completo, despliegue en servidor de desarrollo y configuración de dominio.", style='List Bullet')

    # Precio
    doc.add_heading('4. INVERSIÓN', level=1)
    price_table = doc.add_table(rows=1, cols=2)
    hdr_cells = price_table.rows[0].cells
    hdr_cells[0].text = 'Concepto'
    hdr_cells[1].text = 'Valor (COP)'
    
    row_cells = price_table.add_row().cells
    row_cells[0].text = 'Desarrollo web integral Artetex (Diseño, IA Lab, SEO)'
    row_cells[1].text = '$600.000'

    doc.add_paragraph("\nVALOR TOTAL: SEISCIENTOS MIL PESOS M/CTE (600.000 COP)")

    # Firma
    doc.add_paragraph("\n\n" + "_" * 40)
    doc.add_paragraph("Jorge Jimenez Calvano\nDesarrollo de Software & Diseño Digital")

    # Guardar
    doc.save('COTIZACION_ARTETEX_2026.docx')
    print("Documento creado exitosamente: COTIZACION_ARTETEX_2026.docx")

if __name__ == "__main__":
    create_quotation()
